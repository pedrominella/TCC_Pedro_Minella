# -*- coding: utf-8 -*-
r"""
gerar_tcc_completo_abnt.py

Este script realiza a formatação completa do TCC do Pedro Minella sob as Normas ABNT:
1. Margens ABNT (Superior 3cm, Esquerda 3cm, Inferior 2cm, Direita 2cm).
2. Tamanho do papel A4.
3. Formatação de parágrafos: Fonte Times New Roman 12pt, espaçamento de 1.5,
   alinhamento justificado e recuo de primeira linha de 1.25 cm (corpo de texto).
4. Formatação de títulos principales e apêndices (completamente IDEMPOTENTE):
   - Heading 1: Caixa alta (UPPERCASE), Negrito, 12pt, alinhado à esquerda.
   - Heading 2: Caixa baixa (Title Case), Negrito, 12pt, alinhado à esquerda.
   - Heading 3: Itálico e Negrito, 12pt, alinhado à esquerda.
5. Correção e congelamento da numeração e dos títulos dos capítulos e apêndices de forma semântica:
   - 1 INTRODUÇÃO
   - 2 REVISÃO BIBLIOGRÁFICA
   - 3 METODOLOGIA
   - 4 RESULTADOS (primeiro capítulo de resultados)
   - 5 CONCLUSÃO (capítulo de conclusão da pesquisa)
   - REFERÊNCIAS
   - APÊNDICE A — ROBUSTEZ COM MODELOS VAR
   - APÊNDICE B — TESTES DE ESTACIONARIEDADE (ADF E KPSS)
6. Substituição de travessões por pontuação adequada (como vírgulas ou parênteses) no corpo do texto.
7. Formatação de todas as 22 tabelas no formato ABNT (sem linhas verticais, apenas linhas horizontais no topo, base do cabeçalho e base da tabela). Fonte interna 10pt, espaçamento simples (1.0), centralizadas.
8. Criação de estilos personalizados para legendas (LegendaFigura, LegendaTabela, LegendaQuadro)
   e inserção de CAMPOS NATIVOS do Word ("Sumário do Word") para as Listas e o Sumário,
   permitindo a atualização automática com alinhamento perfeito e pontilhado:
   - Lista de Figuras: TOC \t "LegendaFigura;1" \h \z
   - Lista de Tabelas: TOC \t "LegendaTabela;1" \h \z
   - Lista de Quadros: TOC \t "LegendaQuadro;1" \h \z
   - Sumário: TOC \o "1-3" \h \z \u
9. Preservação dos gráficos (shapes) e dados originais das tabelas do Word.
"""

import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import re
import shutil
import os

def set_abnt_margins(section):
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Aplica margens internas (padding) às células das tabelas para melhor legibilidade."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_abnt_table_borders(table):
    """Aplica bordas padrão ABNT nas tabelas do Word: sem linhas verticais, apenas horizontais no topo, cabeçalho e base."""
    # Remove as bordas gerais da tabela
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is not None:
        tblPr.remove(tblBorders)
        
    # Aplica as bordas específicas por célula (cabeçalho e base)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            # Remove bordas anteriores da célula
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is not None:
                tcPr.remove(tcBorders)
                
            if r_idx == 0:
                # Primeira linha (Cabeçalho): borda superior e inferior
                borders_xml = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                    f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                    f'  <w:left w:val="none"/>'
                    f'  <w:right w:val="none"/>'
                    f'</w:tcBorders>'
                )
            elif r_idx == len(table.rows) - 1:
                # Última linha (Base): borda inferior
                borders_xml = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'  <w:top w:val="none"/>'
                    f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                    f'  <w:left w:val="none"/>'
                    f'  <w:right w:val="none"/>'
                    f'</w:tcBorders>'
                )
            else:
                # Linhas intermediárias: sem bordas horizontais ou verticais
                borders_xml = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'  <w:top w:val="none"/>'
                    f'  <w:bottom w:val="none"/>'
                    f'  <w:left w:val="none"/>'
                    f'  <w:right w:val="none"/>'
                    f'</w:tcBorders>'
                )
            tcPr.append(borders_xml)

def is_equation(text):
    t = text.strip()
    if not t:
        return False
    # Identifica fórmulas por conter símbolos matemáticos e ter texto curto
    math_syms = ['=', '+', '-', '×', 'Σ', 'α', 'β', 'γ', 'δ', 'θ', 'ε', 'π', 'λ', 'Ψ', 'σ', 'Φ', 'Δ']
    if any(sym in t for sym in math_syms) and len(t) < 120:
        exclude_starts = ('Figura ', 'Tabela ', 'Quadro ', 'Fonte:', 'Nota:', '1 ', '2 ', '3 ', '4 ', '5 ', '6 ', 'APÊNDICE')
        if not t.startswith(exclude_starts):
            return True
    return False

def replace_text_in_runs(p, replacements):
    """Substitui texto nos runs de forma a manter os estilos de negrito/itálico intactos."""
    for run in p.runs:
        for k, v in replacements.items():
            if k in run.text:
                run.text = run.text.replace(k, v)

def add_native_word_field(p, instruction):
    """Insere um campo nativo do Word (como TOC ou lista de figuras) no parágrafo com estilo Normal."""
    p.text = "" # limpa texto estático anterior
    p_el = p._p
    
    instruction_escaped = instruction.replace('"', '&quot;')
    
    # Criamos o elemento w:fldSimple no XML
    fld_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr="{instruction_escaped}">'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'      <w:sz w:val="24"/>' # 12pt
        f'    </w:rPr>'
        f'    <w:t>Clique com o botão direito aqui e selecione "Atualizar Campo" para gerar esta lista automaticamente.</w:t>'
        f'  </w:r>'
        f'</w:fldSimple>'
    )
    fldSimple = parse_xml(fld_xml)
    p_el.append(fldSimple)

def get_or_create_style(doc, name, pt_size=10, italic=False, bold=False):
    """Obtém ou cria um estilo personalizado no documento para legendas e formatação específica."""
    styles = doc.styles
    try:
        style = styles[name]
    except KeyError:
        style = styles.add_style(name, docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(pt_size)
        font.italic = italic
        font.bold = bold
    return style

def main():
    source_file = 'TCC_Pedro_v7_resultados_finais.docx'
    temp_file = 'TCC_Pedro_ABNT_temp.docx'
    
    print(f"Lendo '{source_file}'...")
    shutil.copy(source_file, temp_file)
    doc = docx.Document(temp_file)
    
    # 1. IDENTIFICAÇÃO IDEMPOTENTE E SEGURA DO INÍCIO DO CORPO DO TEXTO
    # Procuramos o parágrafo que inicia com "1 INTRODUÇÃO" e tem estilo Heading 1!
    intro_idx = None
    for idx, p in enumerate(doc.paragraphs):
        text_clean = p.text.strip().upper()
        if p.style.name.startswith('Heading 1') and (text_clean.startswith('1 INTRODUÇÃO') or text_clean.startswith('1 INTRODUCAO')):
            intro_idx = idx
            break
            
    if intro_idx is None:
        # Fallback caso não ache com Heading 1, busca pelo texto no início de parágrafos
        for idx, p in enumerate(doc.paragraphs):
            text_clean = p.text.strip().upper()
            if text_clean.startswith('1 INTRODUÇÃO') or text_clean.startswith('1 INTRODUCAO'):
                intro_idx = idx
                break

    if intro_idx is not None:
        print(f"Encontrou a Introdução real no parágrafo {intro_idx}. Deletando todos os {intro_idx} parágrafos anteriores...")
        # Deletamos todos os parágrafos pré-textuais anteriores de forma limpa e idempotente
        for _ in range(intro_idx):
            p = doc.paragraphs[0]
            p._element.getparent().remove(p._element)
    else:
        print("Erro: Não foi possível localizar o início da '1 INTRODUÇÃO'. A formatação pode falhar.")
        return

    # Agora o parágrafo index 0 é a Introdução real!
    first_p = doc.paragraphs[0]
    
    # 2. Ajustes de Numeração e Texto nos Títulos e Corpo de Texto (renomeação 5 -> 4, 6 -> 5, etc.)
    number_replacements = {
        'Capítulo 5': 'Capítulo 4',
        'capítulo 5': 'capítulo 4',
        'Capítulo 6': 'Capítulo 5',
        'capítulo 6': 'capítulo 5',
        'seção 5.1': 'seção 4.1',
        'Seção 5.1': 'Seção 4.1',
        'seção 5.2': 'seção 4.2',
        'Seção 5.2': 'Seção 4.2',
        'seção 5.3': 'seção 4.3',
        'Seção 5.3': 'Seção 4.3',
        'seção 5.4': 'seção 4.5.1',
        'Seção 5.4': 'Seção 4.5.1',
        'seção 5.5': 'seção 4.4',
        'Seção 5.5': 'Seção 4.4',
        'seção 5.6': 'seção 4.5',
        'Seção 5.6': 'Seção 4.5',
        'seção 5.7': 'seção 4.6',
        'Seção 5.7': 'Seção 4.6',
        'Apêndice B': 'Apêndice A',
        'apêndice B': 'apêndice A',
        'Apêndice C': 'Apêndice B',
        'apêndice C': 'apêndice B',
        'Apêndice D': 'Apêndice C',
        'apêndice D': 'apêndice C',
    }
    
    print("Corrigindo títulos principais e apêndices de forma IDEMPOTENTE com look-ahead semântico...")
    
    # Renomeando títulos e capturando dados
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style_name = p.style.name
        
        # Correção específica de títulos de apêndice nos títulos de figuras e tabelas
        if text.startswith('Figura C.'):
            new_text = text.replace('Figura C.', 'Figura B.')
            replace_text_in_runs(p, {'Figura C.': 'Figura B.'})
            text = new_text
        elif text.startswith('Tabela C.'):
            new_text = text.replace('Tabela C.', 'Tabela B.')
            replace_text_in_runs(p, {'Tabela C.': 'Tabela B.'})
            text = new_text
            
        # Renomeação de capítulos e apêndices nos títulos principais de forma idempotente e baseada em semântica
        if style_name.startswith('Heading 1'):
            text_upper = text.upper()
            
            # Olhamos à frente para descobrir o conteúdo real deste capítulo
            next_p = None
            for j in range(idx + 1, min(idx + 5, len(doc.paragraphs))):
                if doc.paragraphs[j].text.strip():
                    next_p = doc.paragraphs[j]
                    break
            next_text = next_p.text.strip().upper() if next_p else ""
            
            if 'INTRODUÇÃO' in text_upper or 'INTRODUCAO' in text_upper:
                p.text = '1 INTRODUÇÃO'
            elif 'REVISÃO BIBLIOGRÁFICA' in text_upper or 'REVISAO BIBLIOGRAFICA' in text_upper:
                p.text = '2 REVISÃO BIBLIOGRÁFICA'
            elif 'METODOLOGIA' in text_upper:
                p.text = '3 METODOLOGIA'
            elif 'REFERÊNCIAS' in text_upper or 'REFERENCIAS' in text_upper:
                p.text = 'REFERÊNCIAS'
            elif 'APÊNDICE' in text_upper or 'APENDICE' in text_upper or 'APÊNDICE' in next_text or 'APENDICE' in next_text or 'TESTES DE ESTACIONARIEDADE' in text_upper or 'ROBUSTEZ COM MODELOS VAR' in text_upper:
                # O primeiro apêndice de verdade tem o texto de VAR.
                # O segundo apêndice tem o texto de Dickey-Fuller/KPSS.
                if 'VETORIAIS AUTORREGRESSIVOS' in next_text or 'ESTE APÊNDICE DOCUMENTA' in next_text:
                    p.text = 'APÊNDICE A — ROBUSTEZ COM MODELOS VAR'
                else:
                    p.text = 'APÊNDICE B — TESTES DE ESTACIONARIEDADE (ADF E KPSS)'
            elif 'RESULTADOS' in text_upper or 'CONCLUS' in text_upper:
                # Se o texto que segue contém "Este trabalho analisou", é a conclusão.
                # Caso contrário, é o capítulo principal de Resultados.
                if 'ESTE TRABALHO ANALISOU' in next_text or 'OS RESULTADOS SUSTENTAM' in next_text or 'ESTE TRABALHO ANALISA' in next_text or 'LIMITAÇÃO DO TRABALHO' in next_text:
                    p.text = '5 CONCLUSÃO'
                else:
                    p.text = '4 RESULTADOS'
                
        elif style_name.startswith('Heading 2'):
            # Formatação e renumeração de subseções de forma idempotente
            if text.startswith('5.1') or text.startswith('4.1'):
                p.text = '4.1 Repasse do petróleo aos combustíveis'
            elif text.startswith('5.2') or text.startswith('4.2'):
                p.text = '4.2 Efeito dos combustíveis sobre o IPCA Geral'
            elif text.startswith('5.3') or text.startswith('4.3'):
                p.text = '4.3 Efeito dos combustíveis sobre o IPCA Transportes'
            elif text.startswith('5.5') or text.startswith('4.4'):
                p.text = '4.4 O canal dos combustíveis como mediador entre petróleo e inflação'
            elif text.startswith('5.6') or text.startswith('4.5'):
                p.text = '4.5 Exercícios de robustez'
            elif text.startswith('5.7') or text.startswith('4.6'):
                p.text = '4.6 Síntese dos resultados'
                
        elif style_name.startswith('Heading 3'):
            if text.startswith('5.4') or text.startswith('4.5.1'):
                p.text = '4.5.1 O Papel da Política de Preços da Petrobras: Análise por Regimes (Pré e Pós-2016)'
            elif text.startswith('5.6.1') or text.startswith('4.5.2'):
                p.text = '4.5.2 LP-IV com Oil Supply News Shock'
                
        # Substitui referências no corpo do texto
        replace_text_in_runs(p, number_replacements)
        
        # Eliminar travessões (em-dashes) no corpo do texto (exceto equações e legendas)
        if '—' in p.text and not is_equation(p.text) and not p.style.name.startswith('Heading') and not p.text.startswith(('Figura ', 'Tabela ', 'Quadro ', 'APÊNDICE')):
            t_count = p.text.count('—')
            if t_count == 2:
                parts = p.text.split('—')
                new_text = parts[0] + '(' + parts[1].strip() + ') ' + parts[2]
                for r in p.runs:
                    r.text = ''
                p.runs[0].text = new_text
            elif t_count == 1:
                new_text = p.text.replace(' — ', ', ').replace('—', ', ')
                for r in p.runs:
                    r.text = ''
                p.runs[0].text = new_text

    print("Corrigindo tabelas...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text_in_runs(p, number_replacements)
                    if '—' in p.text and not is_equation(p.text):
                        p.text = p.text.replace(' — ', ', ').replace('—', ', ')
    
    # 3. Gerando os Elementos Pré-Textuais sob Norma ABNT
    print("Gerando elementos pré-textuais novos com CAMPOS NATIVOS do Word...")
    
    def add_p(text, size=12, bold=False, space_before=0, space_after=6, align=WD_ALIGN_PARAGRAPH.CENTER, uppercase=False):
        """Helper para inserir parágrafo formatado no início com estilo 'Normal' explícito."""
        p = first_p.insert_paragraph_before(style='Normal')
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = align
        run = p.add_run(text.upper() if uppercase else text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold
        return p
        
    # --- CAPA ---
    add_p("Instituto Brasileiro de Mercado de Capitais - Ibmec-DF", size=12, bold=True, uppercase=True)
    add_p("Curso de Bacharelado em Ciências Econômicas", size=12, bold=True, space_after=48, uppercase=True)
    
    add_p("Pedro Franck Minella", size=12, bold=True, space_before=72, space_after=96, uppercase=True)
    
    add_p("Efeitos dos Choques do Preço do Petróleo sobre a Inflação no Brasil:", size=14, bold=True, space_before=72)
    add_p("Evidências via Combustíveis, Petrobras, VAR e Local Projections entre 2003 e 2026", size=14, bold=True, space_after=144)
    
    add_p("Brasília/DF", size=12, bold=True, space_before=144)
    add_p("2026", size=12, bold=True)
    
    # Quebra de página após a Capa
    p_break = first_p.insert_paragraph_before(style='Normal')
    p_break.add_run().add_break(WD_BREAK.PAGE)
    
    # --- FOLHA DE ROSTO ---
    add_p("Pedro Franck Minella", size=12, bold=True, space_after=48, uppercase=True)
    
    add_p("Efeitos dos Choques do Preço do Petróleo sobre a Inflação no Brasil:", size=14, bold=True, space_before=72)
    add_p("Evidências via Combustíveis, Petrobras, VAR e Local Projections entre 2003 e 2026", size=14, bold=True, space_after=48)
    
    # Nota da Folha de Rosto (Recuo de 8cm, tamanho 10, justificado, espaçamento simples)
    p_note = first_p.insert_paragraph_before(style='Normal')
    p_note.paragraph_format.line_spacing = 1.0
    p_note.paragraph_format.space_before = Pt(24)
    p_note.paragraph_format.space_after = Pt(48)
    p_note.paragraph_format.left_indent = Cm(8.0)
    p_note.paragraph_format.first_line_indent = Cm(0)
    p_note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_note = p_note.add_run(
        "Trabalho de Conclusão de Curso apresentado como requisito para obtenção do título de "
        "Bacharel em Ciências Econômicas pelo Instituto Brasileiro de Mercado de Capitais (Ibmec-DF).\n\n"
        "Orientador: Prof. Silvio Costa\n"
        "Coordenador: Prof. Frederico Dias"
    )
    run_note.font.name = 'Times New Roman'
    run_note.font.size = Pt(10)
    
    add_p("Brasília/DF", size=12, bold=True, space_before=96)
    add_p("2026", size=12, bold=True)
    
    # Quebra de página após a Folha de Rosto
    p_break2 = first_p.insert_paragraph_before(style='Normal')
    p_break2.add_run().add_break(WD_BREAK.PAGE)
    
    # --- RESUMO ---
    p_res_title = first_p.insert_paragraph_before(style='Normal')
    p_res_title.paragraph_format.space_before = Pt(12)
    p_res_title.paragraph_format.space_after = Pt(12)
    p_res_title.paragraph_format.first_line_indent = Cm(0)
    p_res_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_res_title = p_res_title.add_run("RESUMO")
    run_res_title.font.name = 'Times New Roman'
    run_res_title.font.size = Pt(12)
    run_res_title.font.bold = True
    
    p_res_text = first_p.insert_paragraph_before(style='Normal')
    p_res_text.paragraph_format.line_spacing = 1.0  # Espaçamento simples
    p_res_text.paragraph_format.space_before = Pt(0)
    p_res_text.paragraph_format.space_after = Pt(12)
    p_res_text.paragraph_format.first_line_indent = Cm(0)  # Sem recuo de primeira linha
    p_res_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_res_text = p_res_text.add_run(
        "Este trabalho analisa em que medida choques no preço internacional do petróleo são transmitidos "
        "aos preços domésticos dos combustíveis e à inflação brasileira entre 2003 e 2026. A hipótese "
        "central é que o repasse ocorre principalmente por meio dos combustíveis, com maior intensidade "
        "no IPCA Transportes do que no IPCA Geral, e que a política de preços da Petrobras altera o timing "
        "e a magnitude dessa transmissão. A estratégia empírica principal utiliza Local Projections com "
        "respostas acumuladas, controles macroeconômicos, defasagens e inferência HAC/Newey-West. O trabalho "
        "também estima especificações com Brent em reais, modelos por regimes associados à Petrobras, testes "
        "de Wald para diferenças entre períodos e um exercício LP-IV que utiliza a série Oil Supply News Shock "
        "como instrumento externo para o petróleo em reais. Como robustez adicional, são utilizados modelos "
        "VAR com funções impulso-resposta, respostas acumuladas, FEVD e diagnósticos de resíduos. Os resultados "
        "indicam que os choques de petróleo são transmitidos de forma mais clara aos combustíveis, especialmente "
        "ao diesel e à gasolina de refinaria. O efeito sobre o IPCA Geral é limitado, enquanto o IPCA Transportes "
        "apresenta resposta mais forte e consistente. As evidências por regime indicam heterogeneidade no repasse, "
        "e o instrumento externo confirma que os efeitos inflacionários são mais visíveis no curto prazo e nos "
        "componentes ligados a transporte."
    )
    run_res_text.font.name = 'Times New Roman'
    run_res_text.font.size = Pt(12)
    
    p_kw = first_p.insert_paragraph_before(style='Normal')
    p_kw.paragraph_format.line_spacing = 1.0
    p_kw.paragraph_format.space_before = Pt(12)
    p_kw.paragraph_format.space_after = Pt(12)
    p_kw.paragraph_format.first_line_indent = Cm(0)
    p_kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_kw_lbl = p_kw.add_run("Palavras-chave: ")
    run_kw_lbl.font.name = 'Times New Roman'
    run_kw_lbl.font.size = Pt(12)
    run_kw_lbl.font.bold = True
    run_kw_val = p_kw.add_run("petróleo; combustíveis; inflação; Petrobras; Local Projections.")
    run_kw_val.font.name = 'Times New Roman'
    run_kw_val.font.size = Pt(12)
    
    # Quebra de página após o Resumo
    p_break3 = first_p.insert_paragraph_before(style='Normal')
    p_break3.add_run().add_break(WD_BREAK.PAGE)
    
    # --- LISTA DE FIGURAS ---
    add_p("LISTA DE FIGURAS", size=12, bold=True, space_before=12, space_after=24)
    # Adiciona o campo nativo do Word para a Lista de Figuras baseado no estilo personalizado
    p_fig_toc = first_p.insert_paragraph_before(style='Normal')
    p_fig_toc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_native_word_field(p_fig_toc, 'TOC \\t "LegendaFigura;1" \\h \\z')
        
    p_break4 = first_p.insert_paragraph_before(style='Normal')
    p_break4.add_run().add_break(WD_BREAK.PAGE)
    
    # --- LISTA DE TABELAS ---
    add_p("LISTA DE TABELAS", size=12, bold=True, space_before=12, space_after=24)
    # Adiciona o campo nativo do Word para a Lista de Tabelas baseado no estilo personalizado
    p_tbl_toc = first_p.insert_paragraph_before(style='Normal')
    p_tbl_toc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_native_word_field(p_tbl_toc, 'TOC \\t "LegendaTabela;1" \\h \\z')
        
    p_break5 = first_p.insert_paragraph_before(style='Normal')
    p_break5.add_run().add_break(WD_BREAK.PAGE)
    
    # --- LISTA DE QUADROS ---
    add_p("LISTA DE QUADROS", size=12, bold=True, space_before=12, space_after=24)
    # Adiciona o campo nativo do Word para a Lista de Quadros baseado no estilo personalizado
    p_qdr_toc = first_p.insert_paragraph_before(style='Normal')
    p_qdr_toc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_native_word_field(p_qdr_toc, 'TOC \\t "LegendaQuadro;1" \\h \\z')
        
    p_break_qdr = first_p.insert_paragraph_before(style='Normal')
    p_break_qdr.add_run().add_break(WD_BREAK.PAGE)
    
    # --- SUMÁRIO ---
    add_p("SUMÁRIO", size=12, bold=True, space_before=12, space_after=24)
    # Adiciona o campo nativo do Word ("Sumário do Word") com estilo Normal
    p_sum_toc = first_p.insert_paragraph_before(style='Normal')
    p_sum_toc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_native_word_field(p_sum_toc, 'TOC \\o "1-3" \\h \\z \\u')
            
    p_break6 = first_p.insert_paragraph_before(style='Normal')
    p_break6.add_run().add_break(WD_BREAK.PAGE)
    
    # 4. Processando os Parágrafos do Corpo do Texto (Formatando margens, fontes e recuos)
    print("Aplicando formatação ABNT de parágrafos em todo o documento...")
    
    # Definindo margens do documento para cada seção
    for section in doc.sections:
        set_abnt_margins(section)
        
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
            
        style_name = p.style.name
        
        # Heading 1 (Capítulos)
        if style_name.startswith('Heading 1'):
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.italic = False
                
        # Heading 2 (Subcapítulos)
        elif style_name.startswith('Heading 2'):
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.italic = False
                
        # Heading 3 (Seções terciárias)
        elif style_name.startswith('Heading 3'):
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.italic = True
                
        # Legendas de figuras
        elif text.startswith('Figura '):
            p.style = get_or_create_style(doc, 'LegendaFigura', pt_size=10, bold=False)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.italic = False
                if 'Figura ' in run.text:
                    run.font.bold = True
                    
        # Legendas de tabelas
        elif text.startswith('Tabela '):
            p.style = get_or_create_style(doc, 'LegendaTabela', pt_size=10, bold=False)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.italic = False
                if 'Tabela ' in run.text:
                    run.font.bold = True
                    
        # Legendas de quadros
        elif text.startswith('Quadro '):
            p.style = get_or_create_style(doc, 'LegendaQuadro', pt_size=10, bold=False)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.italic = False
                if 'Quadro ' in run.text:
                    run.font.bold = True

        # Fontes e Notas das ilustrações
        elif text.startswith(('Fonte:', 'Nota:')):
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                if any(x in run.text for x in ('Fonte:', 'Nota:')):
                    run.font.bold = True
                
        # Fórmulas e Equações
        elif is_equation(text):
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                
        # Citações Longas Recuadas (Estilos com recuo esquerdo no template)
        elif p.paragraph_format.left_indent and p.paragraph_format.left_indent >= Cm(3.0) and len(text) > 100:
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Cm(4.0)  # Recuo de 4cm padrão ABNT
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)  # Tamanho 10 para citação recuada
                
        # Corpo de texto normal
        else:
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Cm(1.25)  # Recuo de primeira linha 1.25cm
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                
    # 5. Formatação das Tabelas no Padrão ABNT
    print("Aplicando bordas ABNT e formatação interna em todas as 22 tabelas do documento...")
    for idx, table in enumerate(doc.tables):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = True
        set_abnt_table_borders(table)
        
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                        if r_idx == 0:
                            run.font.bold = True

    # 6. Salvar documento formatado
    output_name = 'TCC_Pedro_v7_resultados_finais.docx'
    doc.save(output_name)
    print(f"\nDocumento formatado com sucesso sob a norma ABNT e salvo em '{output_name}'!")
    
    if os.path.exists(temp_file):
        os.remove(temp_file)

if __name__ == '__main__':
    main()
