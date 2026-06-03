# -*- coding: utf-8 -*-
import docx

doc = docx.Document('TCC_Pedro_v7_resultados_finais.docx')
print(f"Total tables in document: {len(doc.tables)}")
for idx, table in enumerate(doc.tables):
    print(f"Table {idx}: {len(table.rows)} rows, {len(table.columns)} columns")
    # print first row to identify
    first_row_text = [cell.text.strip() for cell in table.rows[0].cells]
    print(f"  Header: {first_row_text[:5]}")
