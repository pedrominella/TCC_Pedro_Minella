# -*- coding: utf-8 -*-
import docx

doc = docx.Document('TCC_Pedro_v7_resultados_finais.docx')
with open('inspect_first_100.txt', 'w', encoding='utf-8') as f:
    for idx in range(min(120, len(doc.paragraphs))):
        p = doc.paragraphs[idx]
        f.write(f"Para {idx} ({p.style.name}): '{p.text[:120]}'\n")
