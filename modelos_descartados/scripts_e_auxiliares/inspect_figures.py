# -*- coding: utf-8 -*-
import docx

doc = docx.Document('TCC_Pedro_v7_resultados_finais.docx')
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")

# count inline shapes (images)
inline_shapes = doc.inline_shapes
print(f"Total inline shapes: {len(inline_shapes)}")

# Find all figure placeholders in text
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if '[Inserir' in text or 'Figura ' in text:
        print(f"Para {i} ({p.style.name}): '{text}'")
