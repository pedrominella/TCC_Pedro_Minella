# -*- coding: utf-8 -*-
import docx

doc = docx.Document('TCC_Pedro_v7_resultados_finais.docx')
with open('tcc_paras_340_360.txt', 'w', encoding='utf-8') as f:
    for idx in range(335, min(370, len(doc.paragraphs))):
        p = doc.paragraphs[idx]
        f.write(f"Para {idx} ({p.style.name}): '{p.text}'\n")
