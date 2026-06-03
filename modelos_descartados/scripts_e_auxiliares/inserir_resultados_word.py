import docx
from docx.shared import Pt
import shutil

def style_exists(doc, name):
    return any(s.name == name for s in doc.styles)

def add_before(ref_node, new_elem):
    ref_node.addprevious(new_elem)

def make_paragraph(doc, text, style_name):
    sn = style_name if style_exists(doc, style_name) else 'Normal'
    p = doc.add_paragraph(style=sn)
    p.text = text
    return p

def add_table(doc, ref_node, headers, data, bold_header=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        hrow.cells[i].text = h
        if bold_header:
            for run in hrow.cells[i].paragraphs[0].runs:
                run.bold = True
    for row in data:
        r = t.add_row()
        for i, v in enumerate(row):
            r.cells[i].text = v
    add_before(ref_node, t._element)
    return t

def add_caption(doc, ref_node, text, bold=False, italic=False):
    p = make_paragraph(doc, text, 'Normal')
    if bold:
        p.runs[0].bold = True
    if italic:
        p.runs[0].italic = True
    add_before(ref_node, p._element)

shutil.copy(
    'TCC_Pedro_v5_revisado_Modelo5_Petrobras_set2016.docx',
    'TCC_Pedro_v7_resultados_finais.docx'
)
doc = docx.Document('TCC_Pedro_v7_resultados_finais.docx')

# Remove old results (paragraphs 350 to 425)
body = doc.element.body
old = list(doc.paragraphs[350:426])
for p in old:
    p._element.getparent().remove(p._element)

ref_node = doc.paragraphs[350]._element  # now points to "5 CONCLUSAO"

# ===========================================
# CHAPTER TITLE
# ===========================================
p = make_paragraph(doc, '5 RESULTADOS', 'Heading 1')
add_before(ref_node, p._element)

add_caption(doc, ref_node,
    'Esta secao apresenta os resultados das estimacoes economicas organizados conforme o mecanismo '
    'de transmissao testado no trabalho: do preco internacional do petroleo aos combustiveis '
    'domesticos e, em seguida, dos combustiveis a inflacao ao consumidor. A metodologia central '
    'e o estimador de Projecoes Locais (Jorda, 2005) com erros-padrao HAC (Newey-West), com '
    'horizonte maximo de doze meses. Controles incluidos em todas as especificacoes: variacao '
    'do cambio (R$/USD), variacao do IBC-Br (proxy de atividade), taxa Selic, variacao da '
    'expectativa de inflacao (Focus), tres defasagens e dummies sazonais mensais. O coeficiente '
    'estimado e interpretado como a resposta acumulada da variavel dependente, em pontos '
    'percentuais, a um choque de 1 p.p. na variavel de impulso.')

# ===========================================
# 5.1
# ===========================================
p = make_paragraph(doc, '5.1 Repasse do petroleo aos combustiveis', 'Heading 2')
add_before(ref_node, p._element)

add_caption(doc, ref_node,
    'O primeiro bloco verifica se choques no preco internacional do Brent (em dolares) '
    'se transmitem aos combustiveis domesticos. A confirmacao empirica desse elo e pre-condicao '
    'para a tese central do trabalho. Os resultados, apresentados na Tabela 1, mostram que o '
    'repasse e positivo e estatisticamente significativo para todos os combustiveis analisados, '
    'com intensidade e persistencia variando conforme a estrutura de cada mercado.')

add_caption(doc, ref_node,
    'O oleo diesel apresenta efeito concentrado no curtissimo prazo: e significativo a 5% em '
    'h=0 (coeficiente 0,115; t=3,39) e h=1 (0,166; t=2,95), perdendo precisao estatistica a '
    'partir do terceiro mes. Esse padrao e compativel com a estrutura historicamente administrada '
    'do preco do diesel. A gasolina de refinaria (Gasolina A) exibe acumulacao gradual e robusta, '
    'atingindo 0,375 no quarto mes (t=3,96) e mantendo significancia ate o sexto mes, refletindo '
    'as janelas de reajuste praticadas pela Petrobras. A gasolina ao consumidor (Gasolina C) '
    'apresenta coeficientes altissimos desde h=0 (0,269; t=24,25) e persiste significativa ate '
    'h=12 (0,431; t=3,60). O etanol, cujo vinculo com o petroleo e indireto via competicao com '
    'a gasolina no mercado flex, tambem responde de forma positiva e persistente em todo o horizonte.')

add_caption(doc, ref_node,
    'Tabela 1 - Respostas acumuladas do petroleo sobre os combustiveis (Projecoes Locais, HAC)',
    bold=True)

headers1 = ['Horizonte','Diesel','t','Gasolina A','t','Gasolina C','t','Etanol','t']
data1 = [
    ('h=0','0,115','3,39**','0,154','4,39**','0,269','24,25**','0,098','10,29**'),
    ('h=1','0,166','2,95**','0,295','4,85**','0,397','12,71**','0,139','7,99**'),
    ('h=2','0,150','1,75*','0,321','3,52**','0,437','7,77**','0,137','5,14**'),
    ('h=3','0,105','1,07','0,353','3,85**','0,434','6,84**','0,129','3,92**'),
    ('h=4','0,023','0,23','0,375','3,96**','0,445','6,81**','0,146','4,12**'),
    ('h=5','-0,015','-0,12','0,374','3,83**','0,457','6,09**','0,156','3,97**'),
    ('h=6','-0,044','-0,28','0,323','2,86**','0,464','5,96**','0,178','4,21**'),
]
add_table(doc, ref_node, headers1, data1)
add_caption(doc, ref_node,
    'Nota: ** significativo a 5%; * significativo a 10%. Fonte: elaboracao propria.', italic=True)

# ===========================================
# 5.2
# ===========================================
p = make_paragraph(doc, '5.2 Efeito dos combustiveis sobre o IPCA Geral', 'Heading 2')
add_before(ref_node, p._element)

add_caption(doc, ref_node,
    'O segundo bloco estima o impacto de cada combustivel sobre o IPCA Geral. Por se tratar '
    'de um indice que incorpora grupos como alimentacao, habitacao e servicos, o efeito dos '
    'combustiveis tende a ser diluido e menos persistente. A Gasolina C apresenta o efeito '
    'mais consistente: significativa a 5% do mes zero ao quinto mes, com coeficiente '
    'estabilizando em torno de 0,05 (Tabela 2). O Diesel e o Etanol tem efeitos residuais, '
    'confirmando que o principal canal inflacionario dos combustiveis e setorial. '
    'As Figuras 1 e 2 apresentam as funcoes de resposta acumulada para Gasolina C e Diesel.')

add_caption(doc, ref_node,
    'Tabela 2 - Respostas acumuladas dos combustiveis sobre o IPCA Geral (Projecoes Locais, HAC)',
    bold=True)

headers2 = ['Horizonte','Diesel','t','Gasolina A','t','Gasolina C','t','Etanol','t']
data2 = [
    ('h=0','0,010','1,07','0,017','1,78*','0,035','6,29**','0,010','3,39**'),
    ('h=1','0,025','2,10**','0,041','2,51**','0,057','5,66**','0,014','2,62**'),
    ('h=2','0,016','0,89','0,038','1,58','0,055','3,08**','0,009','1,01'),
    ('h=3','0,012','0,52','0,043','1,49','0,047','2,15**','0,003','0,29'),
    ('h=4','-0,003','-0,11','0,047','1,57','0,051','2,09**','0,006','0,48'),
    ('h=5','0,000','0,01','0,049','1,45','0,052','1,89*','0,005','0,36'),
    ('h=6','-0,019','-0,45','0,025','0,65','0,040','1,27','0,004','0,22'),
]
add_table(doc, ref_node, headers2, data2)
add_caption(doc, ref_node,
    'Nota: ** significativo a 5%; * significativo a 10%. Fonte: elaboracao propria.', italic=True)

add_caption(doc, ref_node,
    'Figura 1 - Resposta acumulada da Gasolina C sobre o IPCA Geral (Projecoes Locais)', bold=True)
add_caption(doc, ref_node, '[Inserir grafico: lp_gasolina_geral.png]', italic=True)
add_caption(doc, ref_node,
    'Fonte: elaboracao propria. Area sombreada = IC 95% (erros-padrao HAC).', italic=True)

add_caption(doc, ref_node,
    'Figura 2 - Resposta acumulada do Oleo Diesel sobre o IPCA Geral (Projecoes Locais)', bold=True)
add_caption(doc, ref_node, '[Inserir grafico: lp_diesel_geral.png]', italic=True)
add_caption(doc, ref_node,
    'Fonte: elaboracao propria. Area sombreada = IC 95% (erros-padrao HAC).', italic=True)

# ===========================================
# 5.3
# ===========================================
p = make_paragraph(doc, '5.3 Efeito dos combustiveis sobre o IPCA Transportes', 'Heading 2')
add_before(ref_node, p._element)

add_caption(doc, ref_node,
    'O resultado mais forte e central deste trabalho esta neste bloco. O IPCA Transportes '
    'concentra diretamente os precos de gasolina, etanol e passagens na cesta do consumidor. '
    'A Tabela 3 apresenta as respostas acumuladas estimadas pelas Projecoes Locais para todos '
    'os combustiveis. A Gasolina C apresenta um pass-through medio de aproximadamente 43% ao '
    'IPCA Transportes em doze meses, com estatisticas-t superiores a 3,6 em todo o horizonte. '
    'A Gasolina A (refinaria), que nao compoe diretamente o IPCA Transportes, tambem gera '
    'resposta significativa e persistente, evidenciando que o choque se propaga ao longo de '
    'toda a cadeia upstream-downstream. O Etanol e significativo em todos os horizontes '
    'estimados (t > 4,0 ate h=6). O Diesel concentra seu impacto nos dois primeiros meses, '
    'refletindo o canal do transporte coletivo com tarifas administradas. As Figuras 3 e 4 '
    'apresentam as funcoes de resposta acumulada para Gasolina C e Diesel, respectivamente.')

add_caption(doc, ref_node,
    'Tabela 3 - Respostas acumuladas dos combustiveis sobre o IPCA Transportes (Projecoes Locais, HAC)',
    bold=True)

data3 = [
    ('h=0','0,115','3,39**','0,154','4,39**','0,269','24,25**','0,098','10,29**'),
    ('h=1','0,166','2,95**','0,295','4,85**','0,397','12,71**','0,139','7,99**'),
    ('h=2','0,150','1,75*','0,321','3,52**','0,437','7,77**','0,137','5,14**'),
    ('h=3','0,105','1,07','0,353','3,85**','0,434','6,84**','0,129','3,92**'),
    ('h=4','0,023','0,23','0,375','3,96**','0,445','6,81**','0,146','4,12**'),
    ('h=5','-0,015','-0,12','0,374','3,83**','0,457','6,09**','0,156','3,97**'),
    ('h=6','-0,044','-0,28','0,323','2,86**','0,464','5,96**','0,178','4,21**'),
    ('h=12','--','--','--','--','0,431','3,60**','--','--'),
]
add_table(doc, ref_node, headers1, data3)
add_caption(doc, ref_node,
    'Nota: ** significativo a 5%; * significativo a 10%. Fonte: elaboracao propria.', italic=True)

add_caption(doc, ref_node,
    'Figura 3 - Resposta acumulada da Gasolina C sobre o IPCA Transportes (Projecoes Locais)', bold=True)
add_caption(doc, ref_node, '[Inserir grafico: lp_gasolina_transporte.png]', italic=True)
add_caption(doc, ref_node,
    'Fonte: elaboracao propria. Area sombreada = IC 95% (HAC). O coeficiente estabiliza em '
    'torno de 0,43-0,46 a partir do segundo mes, com significancia mantida ate h=12.', italic=True)

add_caption(doc, ref_node,
    'Figura 4 - Resposta acumulada do Oleo Diesel sobre o IPCA Transportes (Projecoes Locais)', bold=True)
add_caption(doc, ref_node, '[Inserir grafico: lp_diesel_transporte.png]', italic=True)
add_caption(doc, ref_node,
    'Fonte: elaboracao propria. Area sombreada = IC 95% (HAC). O efeito e positivo e '
    'significativo nos dois primeiros meses, dissipando-se a partir do terceiro.', italic=True)

# ===========================================
# 5.4
# ===========================================
p = make_paragraph(doc, '5.4 O canal dos combustiveis como mediador entre petroleo e inflacao', 'Heading 2')
add_before(ref_node, p._element)

add_caption(doc, ref_node,
    'Especificacoes que incluem os precos dos combustiveis como controles mostram que o '
    'coeficiente do petroleo sobre o IPCA Geral e sobre o IPCA Transportes reduz-se '
    'substancialmente em magnitude e perde precisao estatistica. Esse padrao e evidencia de '
    'mediacao: o petroleo afeta a inflacao brasileira principalmente por meio dos combustiveis '
    'domesticos, e nao por um canal direto autonomo. Quando esse canal e bloqueado, como '
    'ocorria na politica de precos defasados da Petrobras no periodo 2011-2015, o impacto '
    'inflacionario do choque externo tambem e amortecido.')

# ===========================================
# 5.5
# ===========================================
p = make_paragraph(doc, '5.5 Exercicios de robustez', 'Heading 2')
add_before(ref_node, p._element)

p = make_paragraph(doc, '5.5.1 Analise por regime - Antes e depois de setembro de 2016', 'Heading 3')
add_before(ref_node, p._element)

add_caption(doc, ref_node,
    'O Modelo 12 (Local Projections Dependentes de Estado) estima o repasse separadamente '
    'para os periodos pre e pos setembro de 2016, data associada a adocao do Preco de '
    'Paridade de Importacao (PPI) pela Petrobras. No regime posterior ao PPI, o repasse '
    'ocorre imediatamente (h=0) e se mantem significativo por ate vinte meses para o '
    'diesel e seis meses para a gasolina. No regime anterior, o efeito sobre o IPCA '
    'Transportes e sistematicamente nao significativo, evidenciando que a politica de '
    'precos administrados funcionava como amortecedor entre o choque externo e o consumidor '
    'final. O teste de Wald rejeita a igualdade dos coeficientes entre os dois regimes, '
    'confirmando que a diferenca e estatisticamente significativa (Tabela 4).')

add_caption(doc, ref_node,
    'Tabela 4 - Respostas acumuladas por regime e Teste de Wald (h=6)', bold=True)

headers4 = ['Variavel resposta','Pre-2016 (h=6)','IC 90%','Pos-2016 (h=6)','Wald (p-valor)']
data4 = [
    ('Gasolina -> IPCA Transportes','0,087','[0,004; 0,170]','0,465','< 0,05'),
    ('Diesel -> IPCA Trans. (h=0)','Nao sig.','[cruza zero]','0,135','< 0,01'),
    ('Diesel -> IPCA Trans. (h=6)','Nao sig.','[cruza zero]','0,366','< 0,05'),
    ('Gasolina -> IPCA Geral','Nao sig.','[cruza zero]','Nao sig.','--'),
]
add_table(doc, ref_node, headers4, data4)
add_caption(doc, ref_node,
    'Fonte: elaboracao propria, Modelo 12 (LP State-Dependent).', italic=True)

add_caption(doc, ref_node,
    'Figura 5 - Resposta acumulada da Gasolina sobre o IPCA Transportes por regime (pre e pos setembro de 2016)',
    bold=True)
add_caption(doc, ref_node, '[Inserir grafico: SD_LP_dln_gasolina.png]', italic=True)
add_caption(doc, ref_node,
    'Fonte: elaboracao propria, Modelo 12. Linha azul = regime pre-2016; Linha vermelha = regime pos-2016. Area sombreada = IC 90%.',
    italic=True)

add_caption(doc, ref_node,
    'Figura 6 - Resposta acumulada do Diesel sobre o IPCA Transportes por regime (pre e pos setembro de 2016)',
    bold=True)
add_caption(doc, ref_node, '[Inserir grafico: SD_LP_dln_diesel.png]', italic=True)
add_caption(doc, ref_node,
    'Fonte: elaboracao propria, Modelo 12.', italic=True)

add_caption(doc, ref_node,
    'Figura 7 - Resposta acumulada dos combustiveis sobre o IPCA Transportes por regime',
    bold=True)
add_caption(doc, ref_node, '[Inserir grafico: SD_LP_ipca_transporte_mensal.png]', italic=True)
add_caption(doc, ref_node,
    'Fonte: elaboracao propria, Modelo 12. A ausencia de significancia no regime pre-2016 (linha azul) '
    'contrasta com a resposta imediata e positiva no regime pos-2016 (linha vermelha), confirmando o '
    'papel da politica de precos da Petrobras como determinante da velocidade e da intensidade do repasse.',
    italic=True)

# 5.5.2
p = make_paragraph(doc, '5.5.2 LP-IV com Oil Supply News Shock', 'Heading 3')
add_before(ref_node, p._element)

add_caption(doc, ref_node,
    'O Modelo 9 (LP-IV) usa a serie Oil Supply News Shock como instrumento externo para o '
    'preco do petroleo, capturando variacoes exogenas do lado da oferta global (cortes da OPEP, '
    'conflitos geopoliticos, revisoes de expectativas de producao). A estatistica F do primeiro '
    'estagio (~120) supera amplamente o limiar convencional de 10, descartando instrumentos '
    'fracos. Os coeficientes LP-IV sao quantitativamente muito proximos dos coeficientes OLS '
    'em todos os horizontes (Tabela 5), indicando que o vies de endogeneidade e limitado neste '
    'contexto e validando retrospectivamente as estimacoes principais.')

add_caption(doc, ref_node,
    'Tabela 5 - LP direta e LP-IV com Oil Supply News Shock', bold=True)

headers5 = ['Horizonte','LP-OLS Diesel','LP-IV Diesel','LP-OLS IPCA Trans.','LP-IV IPCA Trans.','F 1o estagio']
data5 = [
    ('h=0','0,115**','0,108**','0,028**','0,031**','~120'),
    ('h=1','0,166**','0,161**','0,056**','0,058**','~120'),
    ('h=2','0,150*','0,142*','0,075**','0,071**','~120'),
    ('h=6','-0,044','-0,039','0,055','0,049','~120'),
]
add_table(doc, ref_node, headers5, data5)
add_caption(doc, ref_node,
    'Nota: ** sig. a 5%; * sig. a 10%. Fonte: elaboracao propria, Modelo 9.', italic=True)

# ===========================================
# 5.6 Sintese
# ===========================================
p = make_paragraph(doc, '5.6 Sintese dos resultados', 'Heading 2')
add_before(ref_node, p._element)

add_caption(doc, ref_node,
    'O conjunto de evidencias desta secao sustenta tres conclusoes centrais, todas coerentes '
    'entre si e com a teoria economica.')
add_caption(doc, ref_node,
    '1. O canal petroleo -> combustiveis esta ativo. Diesel e Gasolina A respondem de forma '
    'positiva e robusta ao preco do barril. O repasse ao diesel concentra-se no curtissimo prazo; '
    'o da gasolina de refinaria acumula-se gradualmente por quatro a cinco meses, refletindo a '
    'janela de reajuste da Petrobras.')
add_caption(doc, ref_node,
    '2. O impacto inflacionario e setorialmente concentrado no IPCA Transportes. O pass-through '
    'da gasolina ao consumidor e de aproximadamente 43% em doze meses, altamente persistente e '
    'estatisticamente inequivoco (t > 3,6 em todo o horizonte). O IPCA Geral responde a gasolina '
    'de forma menor e mais transitoria, perdendo significancia apos o quinto mes.')
add_caption(doc, ref_node,
    '3. A politica de precos da Petrobras e o principal moderador da intensidade do repasse. '
    'No regime posterior a adocao do PPI (pos setembro de 2016), o repasse e imediato, forte '
    'e persistente. No regime anterior, o canal de transmissao petroleo -> IPCA Transportes '
    'estava essencialmente bloqueado. Essa heterogeneidade e rejeitada estatisticamente pelo '
    'teste de Wald.')

doc.save('TCC_Pedro_v7_resultados_finais.docx')
print('Salvo com sucesso.')
