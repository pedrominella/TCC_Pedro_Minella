import docx
from docx.shared import Pt

doc = docx.Document('TCC_Pedro_v7_resultados_finais.docx')
paras = doc.paragraphs

# ============================================================
# HELPER: update text of a paragraph preserving style/runs
# ============================================================
def set_text(para, new_text):
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)

# ============================================================
# 1. Fix '3 METODOLOGIA' section: remove references to "capítulo 4"
#    and replace with "capítulo 5" where appropriate.
# ============================================================
for p in paras[190:350]:
    if 'Cap' not in p.text and 'cap' not in p.text:
        continue
    old = p.text
    new = (old
           .replace('Cap\u00edtulo 4', 'Cap\u00edtulo 5')
           .replace('cap\u00edtulo 4', 'cap\u00edtulo 5')
           .replace('se\u00e7\u00e3o 4', 'se\u00e7\u00e3o 5')
           .replace('Se\u00e7\u00e3o 4', 'Se\u00e7\u00e3o 5'))
    if new != old:
        set_text(p, new)
        print(f'Fixed cap ref in para: {new[:80]}')

# ============================================================
# 2. Fix conclusion numbering: '5 CONCLUSÃO' -> '6 CONCLUSÃO'
# ============================================================
for p in paras[400:412]:
    if p.text.strip().startswith('5 CONCL'):
        set_text(p, '6 CONCLUS\u00c3O')
        print('Fixed: 5 -> 6 CONCLUSAO')
        break

# ============================================================
# 3. Rewrite the Conclusion text to reflect exact final results
#    (paragraphs 403-408 in the original)
# ============================================================
new_conclusion_paras = [
    ('Este trabalho analisou em que medida choques no preco internacional do petroleo sao '
     'transmitidos aos precos domesticos dos combustiveis e a inflacao brasileira entre 2003 '
     'e 2026. A hipotese central era que o repasse ocorre principalmente por meio dos '
     'combustiveis, com efeito mais evidente no IPCA Transportes do que no IPCA Geral, e '
     'que a politica de precos da Petrobras altera a intensidade e o timing desse processo.'),
    ('Os resultados sustentam essa hipotese. As Projecoes Locais, estimadas com erros-padrao '
     'robustos do tipo HAC/Newey-West, mostram que choques no preco do petroleo elevam os '
     'precos dos combustiveis domesticos de forma positiva e estatisticamente significativa. '
     'O oleo diesel apresenta repasse concentrado nos dois primeiros meses apos o choque '
     '(coeficientes de 0,115 e 0,166 em h=0 e h=1, respectivamente, significativos a 5%). '
     'A gasolina de refinaria acumula repasse gradual e robusto ate o quarto mes, com '
     'coeficiente de 0,375 (t=3,96). A gasolina ao consumidor e altamente significativa em '
     'todos os horizontes, com pass-through acumulado de aproximadamente 43% em doze meses '
     '(coeficiente de 0,431; t=3,60 em h=12). O etanol responde de forma indireta, via '
     'canal de substituicao com a gasolina, e e significativo em todos os horizontes estimados.'),
    ('No que se refere ao impacto sobre a inflacao, o resultado mais forte aparece no IPCA '
     'Transportes. A gasolina ao consumidor eleva o IPCA Transportes de forma imediata e '
     'persistente, com coeficiente de 0,269 no proprio mes do choque (t=24,25) e de 0,464 '
     'em seis meses (t=5,96). Esse nivel de precisao estatistica e incomum em series '
     'temporais macroeconomicas de amostras finitas. A gasolina de refinaria tambem gera '
     'resposta significativa sobre o IPCA Transportes, confirmando que o choque se propaga '
     'por toda a cadeia upstream-downstream, da refinaria ao consumidor final. O IPCA Geral '
     'responde de forma menor e menos persistente: a gasolina ao consumidor e significativa '
     'sobre o indice agregado nos primeiros cinco meses (coeficiente medio de 0,05), mas '
     'perde precisao estatistica nos horizontes mais longos. O diesel e o etanol nao '
     'apresentam efeito robusto sobre o IPCA Geral em nenhum horizonte.'),
    ('A analise por regimes e um dos achados mais relevantes do trabalho. O Modelo 12 '
     '(Local Projections Dependentes de Estado), com corte em setembro de 2016, mostra que '
     'no regime posterior a adocao do Preco de Paridade de Importacao (PPI), o repasse '
     'ocorre de forma imediata e estatisticamente significativa desde o mes do choque. '
     'No regime anterior, os coeficientes sobre o IPCA Transportes sao sistematicamente '
     'nao significativos em todos os horizontes. O teste de Wald rejeita a hipotese de '
     'igualdade entre os coeficientes dos dois regimes, confirmando que a mudanca '
     'institucional alterou a velocidade e a intensidade do repasse. Esses resultados sao '
     'coerentes com a teoria: politicas de precos administrados funcionam como amortecedores '
     'que bloqueiam a transmissao de choques externos. Sua remocao, como ocorreu em 2016, '
     'restaura e amplifica o canal de transmissao.'),
    ('O exercicio LP-IV com Oil Supply News Shock confirma que os resultados nao refletem '
     'vies de endogeneidade. A estatistica F do primeiro estagio (~120) e muito superior '
     'ao limiar convencional de 10, e os coeficientes LP-IV sao quantitativamente proximos '
     'dos coeficientes OLS em todos os horizontes. Essa evidencia reforca a interpretacao '
     'causal dos resultados principais: e o choque externo de oferta de petroleo que causa '
     'o repasse nos combustiveis e na inflacao de transportes, e nao uma correlacao espuria '
     'com outras variaveis macroeconomicas.'),
    ('A principal limitacao do trabalho e que o periodo amostral (2003-2026) inclui episodios '
     'de grande heterogeneidade macroeconomica: a crise financeira global, a fase de precos '
     'administrados da Petrobras (2011-2015), a pandemia de Covid-19, o ciclo de alta do '
     'petroleo e do cambio (2021-2022) e a nova estrategia comercial da Petrobras (2023). '
     'A analise por regime capta parte dessa heterogeneidade, mas nao esgota todos os '
     'subperiodos relevantes. Alem disso, a amostra mensal, embora relativamente longa, '
     'torna-se curta quando dividida em regimes, o que pode reduzir a precisao das '
     'estimativas nas subamostras. Pesquisas futuras poderiam explorar dados regionais '
     'do IPCA, tecnicas de Local Projections com parametros variando no tempo (TVP-LP) '
     'ou suavizacao bayesiana para lidar com a ruido das estimativas em horizontes longos.'),
    ('Em termos de contribuicao, este trabalho organiza de forma sistematica o canal de '
     'transmissao petroleo-combustiveis-inflacao para o Brasil, utilizando a metodologia '
     'de Projecoes Locais como estrategia principal, confirmando que o repasse existe, '
     'e setorialmente concentrado no IPCA Transportes, e condicionado pelo regime de '
     'precificacao da Petrobras. Esses resultados sao robustos a diferentes especificacoes '
     'de controles, a identificacao via instrumento externo e a analise por regime '
     'institucional, fornecendo uma base empirica solida para a analise das politicas '
     'de combustiveis e seus efeitos sobre a inflacao brasileira.'),
]

# Find where the conclusion paragraphs are (after the heading)
concl_start = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '6 CONCLUS\u00c3O':
        concl_start = i + 1
        break

if concl_start is not None:
    # Remove old conclusion paragraphs (find end = REFERENCIAS)
    ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if 'REFER\u00caNCIAS' in p.text or 'REFER' in p.text[:10]:
            ref_idx = i
            break

    if ref_idx:
        old_concl = list(doc.paragraphs[concl_start:ref_idx])
        for p in old_concl:
            p._element.getparent().remove(p._element)
        print(f'Removed {len(old_concl)} old conclusion paragraphs')

    # Re-find reference node
    ref_node = None
    for p in doc.paragraphs:
        if 'REFER\u00caNCIAS' in p.text or 'REFER' in p.text[:10]:
            ref_node = p._element
            break

    if ref_node:
        for txt in new_conclusion_paras:
            np = doc.add_paragraph(style='Normal')
            np.text = txt
            ref_node.addprevious(np._element)
        print('New conclusion inserted.')

# ============================================================
# 4. Fix Appendix B: rename to reflect it's now incorporated
#    into Chapter 5 (robustness). Keep as-is but update title.
# ============================================================
for p in doc.paragraphs:
    if 'AP\u00caNDICE B' in p.text or 'AP\u00caNDICE B' in p.text:
        old = p.text
        new = old.replace(
            'AP\u00caNDICE B - RESULTADOS ATUALIZADOS COM BRENT EM REAIS, REGIMES PETROBRAS E CHOQUES EXTERNOS DE OFERTA',
            'AP\u00caNDICE A - RESULTADOS COMPLEMENTARES: BRENT EM REAIS'
        )
        if new != old:
            set_text(p, new)
            print('Fixed Appendix B title')
        break

for p in doc.paragraphs:
    if 'AP\u00caNDICE C' in p.text or 'APENDICE C' in p.text:
        old = p.text
        new = old.replace('AP\u00caNDICE C', 'AP\u00caNDICE B')
        if new != old:
            set_text(p, new)
            print('Fixed Appendix C -> B')
        break

doc.save('TCC_Pedro_v7_resultados_finais.docx')
print('Documento final salvo com sucesso.')
